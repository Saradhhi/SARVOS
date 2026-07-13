"""
DocumentAgent -- read, extract, search, and summarize PDF, Word, Excel, and
plain-text documents.

All operations are SAFE: reading changes nothing. The risk here is
disclosure, not destruction, so the protection that matters is the sandbox
(document_config.DOCUMENTS_DIR) rather than a confirmation prompt. "Are you
sure you want to read the file you just named?" is theatre; refusing to read
outside the documents directory is not.

Extraction libraries are imported lazily, inside the methods that need them,
so a missing optional dependency degrades to a clear message for that one
format rather than taking down the agent.
"""

from __future__ import annotations

from pathlib import Path

from agents import document_config as config
from agents.automation import resolve_safe_path
from agents.base import BaseAgent
from agents.document_intent import Operation, classify
from core.schemas import AgentName, AgentResult, Task
from llm.client import LLMUnavailable, get_llm_client

SUPPORTED = {".pdf", ".docx", ".xlsx", ".txt", ".md", ".csv"}

_SUMMARY_SYSTEM_PROMPT = (
    "You are summarizing a document for someone who has not read it. Be "
    "accurate and concise. Do not invent details that are not in the text. "
    "If the text appears truncated or incomplete, say so plainly rather "
    "than summarizing as though you saw the whole thing."
)

SEARCH_CONTEXT_CHARS = 80


class DocumentAgent(BaseAgent):
    name = AgentName.DOCUMENT

    def handle(self, task: Task) -> AgentResult:
        intent = classify(task.instruction)
        if intent.operation == Operation.LIST:
            return self._list(task)
        if intent.operation == Operation.UNKNOWN:
            return self._fail(
                task,
                f"I couldn't work out a document action from: "
                f"'{task.instruction}'. Try 'list documents', 'read "
                f"resume.pdf', 'summarize contract.docx', or 'search "
                f"resume.pdf for python'.",
            )

        resolved = self._resolve(task, intent.filename)
        if isinstance(resolved, AgentResult):
            return resolved
        path = resolved

        if intent.operation == Operation.READ:
            return self._read(task, path)
        if intent.operation == Operation.SUMMARIZE:
            return self._summarize(task, path)
        if intent.operation == Operation.SEARCH:
            return self._search(task, path, intent.query)
        return self._fail(task, f"Unsupported document operation: {intent.operation}")

    # ---- helpers ---------------------------------------------------------

    def _fail(self, task: Task, msg: str, error: str = "document_bad_request") -> AgentResult:
        return AgentResult(task_id=task.task_id, agent=self.name, success=False,
                           output=msg, error=error)

    def _resolve(self, task: Task, filename: str):
        """Sandboxed path resolution. Reading arbitrary files is how
        'read ../../.ssh/id_rsa' becomes a capability."""
        try:
            path = resolve_safe_path(filename, workspace_root=config.DOCUMENTS_DIR)
        except Exception as e:
            return self._fail(
                task,
                f"Refusing to read '{filename}': {e}. Documents must be inside "
                f"{config.DOCUMENTS_DIR}.",
                error="unsafe_document_path",
            )
        if not path.is_file():
            return self._fail(
                task, f"'{filename}' isn't in {config.DOCUMENTS_DIR}.",
                error="document_not_found",
            )
        if path.suffix.lower() not in SUPPORTED:
            return self._fail(
                task,
                f"I can't read '{path.suffix}' files. Supported: "
                f"{', '.join(sorted(SUPPORTED))}.",
                error="unsupported_format",
            )
        if path.stat().st_size > config.MAX_FILE_BYTES:
            return self._fail(
                task,
                f"'{path.name}' is too large ({path.stat().st_size:,} bytes) to "
                f"handle honestly without chunking, which isn't built yet.",
                error="document_too_large",
            )
        return path

    # ---- extraction ------------------------------------------------------

    def _extract(self, path: Path) -> str:
        """Returns the document's full text. Raises on failure -- callers
        translate that into a clear message rather than a partial result."""
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            return "\n".join((p.extract_text() or "") for p in reader.pages)
        if suffix == ".docx":
            import docx
            doc = docx.Document(str(path))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.append("\t".join(c.text for c in row.cells))
            return "\n".join(parts)
        if suffix == ".xlsx":
            import openpyxl
            wb = openpyxl.load_workbook(str(path), data_only=True)
            lines = []
            for ws in wb.worksheets:
                lines.append(f"[sheet: {ws.title}]")
                for row in ws.iter_rows(values_only=True):
                    if any(c is not None for c in row):
                        lines.append("\t".join("" if c is None else str(c) for c in row))
            return "\n".join(lines)
        return path.read_text(encoding="utf-8", errors="replace")

    def _extract_or_fail(self, task: Task, path: Path):
        try:
            text = self._extract(path)
        except ImportError as e:
            return self._fail(
                task, f"Can't read '{path.suffix}' files -- a library is missing: {e}",
                error="missing_library",
            )
        except Exception as e:
            return self._fail(task, f"Couldn't read '{path.name}': {e}",
                              error="extraction_failed")
        if not text.strip():
            return self._fail(
                task,
                f"'{path.name}' has no extractable text. If it's a scanned PDF, "
                f"it's images of text, and OCR isn't built yet.",
                error="no_text_found",
            )
        return text

    # ---- operations ------------------------------------------------------

    def _list(self, task: Task) -> AgentResult:
        root = Path(config.DOCUMENTS_DIR)
        if not root.is_dir():
            return AgentResult(
                task_id=task.task_id, agent=self.name, success=True,
                output=f"No documents directory yet. Create {config.DOCUMENTS_DIR} "
                       f"and put files in it.",
                data={"documents": []},
            )
        files = sorted(
            p.name for p in root.rglob("*")
            if p.is_file() and p.suffix.lower() in SUPPORTED
        )
        if not files:
            return AgentResult(
                task_id=task.task_id, agent=self.name, success=True,
                output=f"No readable documents in {config.DOCUMENTS_DIR}.",
                data={"documents": []},
            )
        listing = "\n".join(f"  {f}" for f in files)
        return AgentResult(
            task_id=task.task_id, agent=self.name, success=True,
            output=f"{len(files)} document(s):\n{listing}",
            data={"documents": files},
        )

    def _read(self, task: Task, path: Path) -> AgentResult:
        text = self._extract_or_fail(task, path)
        if isinstance(text, AgentResult):
            return text

        full_length = len(text)
        shown = text[: config.MAX_TEXT_LENGTH]
        truncated = full_length > config.MAX_TEXT_LENGTH
        # Truncation is always announced. Silently showing part of a document
        # is how someone concludes a clause isn't there when it is.
        suffix = (
            f"\n\n[Showing {len(shown):,} of {full_length:,} characters. "
            f"Use 'search {path.name} for <text>' to find specific content.]"
            if truncated else ""
        )
        return AgentResult(
            task_id=task.task_id, agent=self.name, success=True,
            output=f"{path.name}:\n\n{shown}{suffix}",
            data={"file": path.name, "chars": full_length, "truncated": truncated},
        )

    def _search(self, task: Task, path: Path, query: str) -> AgentResult:
        text = self._extract_or_fail(task, path)
        if isinstance(text, AgentResult):
            return text

        needle = query.lower()
        haystack = text.lower()
        hits = []
        start = 0
        while True:
            i = haystack.find(needle, start)
            if i == -1:
                break
            a = max(0, i - SEARCH_CONTEXT_CHARS)
            b = min(len(text), i + len(query) + SEARCH_CONTEXT_CHARS)
            snippet = text[a:b].replace("\n", " ").strip()
            hits.append(snippet)
            start = i + len(needle)
            if len(hits) >= 10:
                break

        if not hits:
            return AgentResult(
                task_id=task.task_id, agent=self.name, success=True,
                output=f"'{query}' does not appear in {path.name}.",
                data={"file": path.name, "query": query, "matches": 0},
            )
        listing = "\n".join(f"  ...{h}..." for h in hits)
        return AgentResult(
            task_id=task.task_id, agent=self.name, success=True,
            output=f"{len(hits)} match(es) for '{query}' in {path.name}:\n{listing}",
            data={"file": path.name, "query": query, "matches": len(hits),
                  "snippets": hits},
        )

    def _summarize(self, task: Task, path: Path) -> AgentResult:
        """The one operation that involves the LLM -- and therefore the one
        that can be confidently wrong.

        Two protections, both learned the hard way in this project:
        1. Truncation is never silent. A model handed the first 12k chars of
           a 90k-char contract will summarize those 12k with total confidence
           and no indication anything is missing.
        2. The output is explicitly labelled as a model's summary, not the
           document. The person should never be unsure whether they are
           reading the contract or a paraphrase of part of it.
        """
        text = self._extract_or_fail(task, path)
        if isinstance(text, AgentResult):
            return text

        full_length = len(text)
        sent = text[: config.MAX_SUMMARY_CHARS]
        truncated = full_length > config.MAX_SUMMARY_CHARS

        prompt = f"Summarize this document.\n\n--- BEGIN ---\n{sent}\n--- END ---"
        if truncated:
            prompt = (
                f"NOTE: this is only the first {len(sent):,} characters of a "
                f"{full_length:,}-character document. Summarize only what you "
                f"can see, and say that it is partial.\n\n{prompt}"
            )

        try:
            summary = get_llm_client().generate(prompt, system=_SUMMARY_SYSTEM_PROMPT)
        except LLMUnavailable as e:
            return self._fail(
                task, f"Can't summarize -- the LLM isn't available: {e}",
                error="llm_unavailable",
            )

        header = f"Summary of {path.name} (generated by the local model, not the document itself)"
        warning = (
            f"\n\n[WARNING: only the first {len(sent):,} of {full_length:,} "
            f"characters were read. This summary cannot describe the rest.]"
            if truncated else ""
        )
        return AgentResult(
            task_id=task.task_id, agent=self.name, success=True,
            output=f"{header}:\n\n{summary.strip()}{warning}",
            data={"file": path.name, "truncated": truncated, "chars_read": len(sent),
                  "chars_total": full_length},
        )
