"""
Real tests against REAL documents, built here with known contents.

Unlike most agents in this project, this one is fully verifiable in the
sandbox: a PDF, Word file, and spreadsheet can be created with contents I
control, then extracted and checked against ground truth. No fake backend,
no monkeypatched library -- the actual pypdf/python-docx/openpyxl code runs.
"""

import pytest

from agents.document import DocumentAgent
from core.schemas import AgentName, Task
from memory.engine import MemoryEngine
from memory.store import Store


def _task(instruction: str) -> Task:
    return Task(parent_request_id="r1", agent=AgentName.DOCUMENT, instruction=instruction)


class _FakeLLM:
    def __init__(self, response="A short summary."):
        self.response = response
        self.prompts = []

    def generate(self, prompt, system=None):
        self.prompts.append(prompt)
        return self.response


@pytest.fixture
def docs(tmp_path, monkeypatch):
    d = tmp_path / "documents"
    d.mkdir()
    monkeypatch.setattr("agents.document_config.DOCUMENTS_DIR", str(d))

    # A real PDF. reportlab is a TEST-ONLY dependency, used to build a file
    # whose contents we chose so extraction can be checked against ground
    # truth. If it's missing, skip only the tests that need a PDF rather than
    # erroring every test in this module -- most of them never touch one.
    try:
        from reportlab.pdfgen import canvas
        c = canvas.Canvas(str(d / "resume.pdf"))
        c.drawString(100, 750, "Saradhi - Senior Engineer")
        c.drawString(100, 730, "Skills: Python, FastAPI, Playwright")
        c.showPage()
        c.save()
    except ImportError:
        pass  # tests needing resume.pdf will skip via the needs_pdf fixture

    # A real Word document, with a table.
    import docx
    doc = docx.Document()
    doc.add_heading("Job Description", 0)
    doc.add_paragraph("We need a Python engineer with FastAPI experience.")
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Salary"
    t.rows[0].cells[1].text = "120000"
    doc.save(str(d / "jd.docx"))

    # A real spreadsheet.
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Scores"
    ws["A1"], ws["B1"] = "Name", "Score"
    ws["A2"], ws["B2"] = "Saradhi", 95
    wb.save(str(d / "scores.xlsx"))

    (d / "notes.txt").write_text("Remember to negotiate the salary.")
    yield d


@pytest.fixture
def agent(tmp_path, docs):
    memory = MemoryEngine(store=Store(tmp_path / "d.db"))
    yield DocumentAgent(memory)


@pytest.fixture
def needs_pdf(docs):
    """Skip cleanly when reportlab (test-only) isn't installed, rather than
    failing a test about SARVOS because of a missing test helper."""
    if not (docs / "resume.pdf").is_file():
        pytest.skip("reportlab not installed -- can't build a real PDF to read")


# ---- Real extraction, checked against contents I wrote -------------------

def test_reads_a_real_pdf(needs_pdf, agent):
    r = agent.handle(_task("read resume.pdf"))
    assert r.success, r.output
    assert "Saradhi - Senior Engineer" in r.output
    assert "FastAPI" in r.output


def test_reads_a_real_docx_including_tables(agent):
    r = agent.handle(_task("read jd.docx"))
    assert r.success
    assert "Job Description" in r.output
    assert "Python engineer" in r.output
    assert "120000" in r.output, "table cells must be extracted too"


def test_reads_a_real_xlsx(agent):
    r = agent.handle(_task("read scores.xlsx"))
    assert r.success
    assert "Saradhi" in r.output
    assert "95" in r.output
    assert "Scores" in r.output  # sheet name


def test_reads_plain_text(agent):
    r = agent.handle(_task("read notes.txt"))
    assert r.success
    assert "negotiate the salary" in r.output


def test_list_documents(agent, docs):
    r = agent.handle(_task("list documents"))
    assert r.success
    expected = {"jd.docx", "scores.xlsx", "notes.txt"}
    if (docs / "resume.pdf").is_file():
        expected.add("resume.pdf")
    assert set(r.data["documents"]) == expected


# ---- Search --------------------------------------------------------------

def test_search_finds_real_text_with_context(needs_pdf, agent):
    r = agent.handle(_task("search resume.pdf for FastAPI"))
    assert r.success
    assert r.data["matches"] >= 1
    assert "FastAPI" in r.data["snippets"][0]


def test_search_is_case_insensitive(needs_pdf, agent):
    assert agent.handle(_task("search resume.pdf for fastapi")).data["matches"] >= 1


def test_search_reports_absence_honestly(needs_pdf, agent):
    r = agent.handle(_task("search resume.pdf for kubernetes"))
    assert r.success
    assert r.data["matches"] == 0
    assert "does not appear" in r.output


def test_find_in_variant(agent):
    r = agent.handle(_task('find "salary" in notes.txt'))
    assert r.success
    assert r.data["matches"] == 1


# ---- Sandboxing ----------------------------------------------------------

def test_refuses_a_path_outside_the_documents_dir(agent):
    """Reading arbitrary files is how 'read ../../.ssh/id_rsa' becomes a
    capability."""
    r = agent.handle(_task("read ../../../etc/passwd.txt"))
    assert not r.success
    assert r.error in {"unsafe_document_path", "document_not_found"}


def test_missing_file_fails_clearly(agent):
    r = agent.handle(_task("read nonexistent.pdf"))
    assert not r.success
    assert r.error == "document_not_found"


def test_unsupported_format_refused(agent, docs):
    (docs / "thing.exe").write_bytes(b"MZ")
    r = agent.handle(_task("read thing.exe"))
    assert not r.success
    assert r.error == "unsupported_format"


def test_oversized_file_refused(agent, docs, monkeypatch):
    monkeypatch.setattr("agents.document_config.MAX_FILE_BYTES", 10)
    r = agent.handle(_task("read notes.txt"))
    assert not r.success
    assert r.error == "document_too_large"


def test_empty_text_reported_not_summarized(agent, docs):
    (docs / "blank.txt").write_text("   \n  ")
    r = agent.handle(_task("read blank.txt"))
    assert not r.success
    assert r.error == "no_text_found"


# ---- Truncation must never be silent -------------------------------------

def test_read_announces_truncation(agent, docs, monkeypatch):
    """Silently showing part of a document is how someone concludes a clause
    isn't there when it is."""
    (docs / "long.txt").write_text("x" * 500)
    monkeypatch.setattr("agents.document_config.MAX_TEXT_LENGTH", 100)
    r = agent.handle(_task("read long.txt"))
    assert r.success
    assert r.data["truncated"] is True
    assert "100 of 500" in r.output.replace(",", "")


# ---- Summarize: the one operation that can be confidently wrong ----------

def test_summarize_sends_the_real_document_text(needs_pdf, agent, monkeypatch):
    fake = _FakeLLM("Saradhi is a senior engineer.")
    monkeypatch.setattr("agents.document.get_llm_client", lambda: fake)
    r = agent.handle(_task("summarize resume.pdf"))
    assert r.success
    assert "Saradhi - Senior Engineer" in fake.prompts[0], "LLM must see real text"
    assert "Saradhi is a senior engineer." in r.output


def test_summary_is_labelled_as_a_summary_not_the_document(agent, monkeypatch):
    """The person must never be unsure whether they're reading the contract
    or a model's paraphrase of part of it."""
    monkeypatch.setattr("agents.document.get_llm_client", lambda: _FakeLLM())
    r = agent.handle(_task("summarize jd.docx"))
    assert "generated by the local model, not the document itself" in r.output


def test_summarize_warns_loudly_when_truncated(agent, docs, monkeypatch):
    """A model handed the first 12k chars of a 90k-char contract will
    summarize those 12k with total confidence and no sign anything is
    missing. The truncation must be visible to the person AND stated to the
    model."""
    (docs / "long.txt").write_text("y" * 5000)
    monkeypatch.setattr("agents.document_config.MAX_SUMMARY_CHARS", 100)
    fake = _FakeLLM("It's a lot of y's.")
    monkeypatch.setattr("agents.document.get_llm_client", lambda: fake)

    r = agent.handle(_task("summarize long.txt"))
    assert r.success
    assert r.data["truncated"] is True
    assert "WARNING" in r.output
    assert "cannot describe the rest" in r.output
    # The model was told, too -- not left to guess.
    assert "only the first" in fake.prompts[0]


def test_summarize_handles_llm_unavailable(needs_pdf, agent, monkeypatch):
    from llm.client import LLMUnavailable

    class Dead:
        def generate(self, prompt, system=None):
            raise LLMUnavailable("Ollama isn't running")

    monkeypatch.setattr("agents.document.get_llm_client", lambda: Dead())
    r = agent.handle(_task("summarize resume.pdf"))
    assert not r.success
    assert r.error == "llm_unavailable"


def test_unrecognized_instruction_gives_helpful_message(agent):
    r = agent.handle(_task("do something documenty but vague"))
    assert not r.success
    assert "list documents" in r.output
